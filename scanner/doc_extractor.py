"""
Extract text thuần từ tài liệu upload (.pdf, .docx, .md, .txt) để đưa vào AI review.

Chạy hoàn toàn local (pypdf + python-docx) — không gửi file ra ngoài, giữ đúng
nguyên tắc "Bảo mật & Privacy" của VulnGuard (xem README).
"""
import logging
import os

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".md", ".txt"}

# Giới hạn ký tự đưa vào prompt AI — tránh vượt context window của model local.
# scanner/doc_reviewer.py sẽ tự cắt theo giới hạn riêng khi build prompt,
# đây chỉ là giới hạn "an toàn" tránh đọc file quá khổng lồ vô ích.
MAX_EXTRACT_CHARS = 200_000


def extract_text(file_path: str, file_ext: str) -> tuple[str, str | None]:
    """Trả về (text, error). text rỗng nếu lỗi, error mô tả lý do nếu có."""
    ext = (file_ext or "").lower()
    try:
        if ext == ".pdf":
            text = _extract_pdf(file_path)
        elif ext == ".docx":
            text = _extract_docx(file_path)
        elif ext in (".md", ".txt"):
            text = _extract_plain(file_path)
        else:
            return "", f"Định dạng file '{ext}' chưa được hỗ trợ — chỉ hỗ trợ .pdf, .docx, .md, .txt"

        text = text.strip()
        if not text:
            return "", "Không trích xuất được nội dung text từ file (file rỗng hoặc là ảnh scan không có OCR)"

        if len(text) > MAX_EXTRACT_CHARS:
            text = text[:MAX_EXTRACT_CHARS]

        return text, None
    except Exception as e:
        logger.error(f"Extract text failed for {file_path}: {e}", exc_info=True)
        return "", f"Lỗi khi đọc file: {e}"


def _extract_pdf(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:
            logger.warning(f"Lỗi extract page {i} của {file_path}: {e}")
    return "\n".join(parts)


def _extract_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text]

    # Lấy luôn nội dung table — SRS/API spec thường mô tả field/quyền trong bảng
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_plain(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def is_supported(filename: str) -> bool:
    _, ext = os.path.splitext(filename.lower())
    return ext in SUPPORTED_EXTS
